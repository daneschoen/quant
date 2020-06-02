import os,sys
import itertools

import numpy as np
#import matplotlib.pyplot as plt

import nltk

from io import StringIO
import docx



"""
Frequency counts of keywords in in document.
Then produces histogram using matplotlib and d3.js

Requires at least: Python 2.7, matplotlib, numpy
Document must be utf-8 or ascii encoded TEXT document

To run:
$ python frequency_words.py frequency_words_doc0.txt   # small document
$ python frequency_words.py frequency_words_doc1.txt   # larger document
$ python frequency_words.py                            # runs an internal test document

Open frequency_d3.html to see d3 visualization
"""


SKIP_SEPARATORS = ['. ', ' .', ',', ';', ':', '"', "'", '`', '~', '*', '!', '?', '(', ')', '{', '}', '<', '>',
                   '+', '=', '|', '\\', '/', '-', '—']
SKIP_SEPARATORS_UNI = SKIP_SEPARATORS + [u'\u2022',u'\u2014']  # the bullet symbol, more can be added ...
SKIP_WORDS = SKIP_SEPARATORS_UNI + ["the", "a", "an", "of", "for", "and", "or", "by", "to", "in", "with", "is", "am", "are", "aren't",
                                    "this", "that", "it", "be", "you", "him", "her", "he", "she", "them", "they" ]


def get_freq(words):
    freq = {}
    for word in words:
        freq[word] = freq.get(word, 0) + 1
    return freq


def get_freq_sorted(words):
    freq = get_freq(words)
    return sorted(freq.items(), key=lambda item: item[1], reverse=True)


def freq_defaultdict(words):
    from collections import Counter
    freq_counter = Counter(words.split())
    return dict(freq_counter)


def get_words(**kwargs):
    _doc = kwargs['doc'].strip()
    _skip_words = []
    if 'skip_words' in kwargs:
        _skip_words = kwargs['skip_words']

    doc_lower = _doc.lower()
    for s in SKIP_SEPARATORS_UNI:
        doc_lower = doc_lower.replace(s, ' ')
    doc_lower_split = doc_lower.split()

    if _skip_words:
        words = [w for w in doc_lower_split if w not in _skip_words]
    else:
        words = doc_lower_split
    return words

"""
def get_words_nltk(**kwargs):
    #
    # Different definition of 'word'; eg, framework.python is one word
    #
    _doc = kwargs['doc']
    _skip_words = []
    if 'skip_words' in kwargs:
        _skip_words = kwargs['skip_words']

    words = []
    for s in nltk.sent_tokenize(_doc.lower()):
        for t in nltk.word_tokenize(s):
            if t in SKIP_SEPARATORS_UNI:
                continue
            if t == '.':
                continue
            words.append(t)

    if _skip_words:
        words = [w for w in words if w not in _skip_words]

    return words
"""

def render_histogram_matplot(freq_sorted, save=True, show=False):
    keys = [kv[0] for kv in freq_sorted]
    freqs = [kv[1] for kv in freq_sorted]

    N = len(keys)
    ind = np.arange(N)    # the x locations for the groups
    width = 0.20

    p1 = plt.bar(ind, freqs, width, color='b')

    plt.ylabel('Word Frequency')
    plt.title('Word Frequency of Document')
    plt.xticks(ind+width/2, keys, rotation=55, fontsize=8)
    if show:
        plt.show()
    if save:
        pass


def output_tsv(freq_sorted, outfile_name):
    try:
        with open(outfile_name, "w") as f:
          f.write("word\tfrequency\n".encode('utf-8'))
          for kv in freq_sorted:
            f.write(kv[0].encode('utf-8') + "\t" + str(kv[1]).encode('utf-8') + "\n")

    except IOError:
        print('Unable to write data file')


def output_csv(freq_sorted, outfile_name):
    try:
        with open(outfile_name, "w") as f:
          f.write("word,frequency\n".encode('utf-8'))
          for kv in freq_sorted:
            f.write(kv[0].encode('utf-8') + "," + str(kv[1]).encode('utf-8') + "\n")

    except IOError:
        print('Unable to write data file')



def get_freq_sorted_json(freq_sorted):
    # [(u'python', 3), (u'django', 3), (u'java', 1), (u'ab\u0107', 1), (u'scraping', 1), (u'framework', 1)]
    # [{"word":"Python","frequency":5},{"word":"Django","frequency":"3"},{"word":"scraping","frequency":2}, {"word":"Java","frequency":1}, {"word":"framework","frequency":1}]
    freq_sorted_json = []
    for x in freq_sorted:
        freq_sorted_json.append({"word":x[0].capitalize(),"frequency":x[1]})

    return freq_sorted_json


def analyze_frequency_words(doc_text):
    #doc_text = unicode(doc_text)

    words = get_words(doc=doc_text, skip_words=SKIP_WORDS)
    # words = get_words_nltk(doc=doc_text, skip_words=SKIP_WORDS)

    freq_sorted = get_freq_sorted(words)
    freq_sorted_json = get_freq_sorted_json(freq_sorted)

    #freq_sorted_json = [{"word":"Python","frequency":5},{"word":"Django","frequency":"3"},{"word":"scraping","frequency":2}, {"word":"Java","frequency":1}, {"word":"framework","frequency":1}]
    return freq_sorted_json



def get_uploaded_file_txt(file_path_name):
    #loc = ""   # '/opt/www/holivue_home/holivue_project/media/documents/'
    doc_text=" "
    try:
      with open(file_path_name, 'rb') as f:
        #doc_text = f.read()
        doc_text = unicode(f.read(), "UTF-8")

    except:
      #print 'Unable to read input file ' + filename
      #sys.exit()
      return None

    else:
      return doc_text.strip()

def get_uploaded_file_docx(file_path_name):
    doc_text=" "
    try:
      #document = docx.Document(filename)

      with open(file_path_name) as f :
      #with open(loc + filename, 'rb') as f :
        #doc_text = f.read()
        #doc_text = unicode(f.read(), "UTF-8")

        document = docx.Document(f)

        #source_stream =  StringIO(f.read())
        #document = docx.Document(source_stream)

        doc_text = '\n\n'.join([
            #paragraph.text.encode('utf-8') for paragraph in document.paragraphs
            paragraph.text for paragraph in document.paragraphs
        ])

    except:
      #print 'Unable to read input file ' + filename
      #sys.exit()
      return None

    else:
      return doc_text.strip()


def get_uploaded_file(filename):
    doc_text=None

    fileName, fileExtension = os.path.splitext(filename)
    if fileExtension == ".docx":
        doc_text = get_uploaded_file_docx(filename)
    elif fileExtension == ".txt":
        doc_text = get_uploaded_file_txt(filename)
    else:
        doc_text = get_uploaded_file_txt(filename)

    return doc_text



if __name__ == '__main__':

    if len(sys.argv) > 1:
         doc_text = get_uploaded_file(sys.argv[1])
         print("Content of file document:")
    else:
        doc_text = u'\u2022 Python, Django, Django AB\u0107, Python; python. django; scraping python. the Java framework'
        print("Using test document:")

    print(doc_text)
    if doc_text and len(doc_text) > 1:
        print(analyze_frequency_words(doc_text))

    """
    Two different ways to run this, custom parser or use nltk
    """
    """
    words = get_words(doc=docu, skip_words=SKIP_WORDS)
    #words = get_words_nltk(doc=docu, skip_words=SKIP_WORDS)

    freq_sorted = get_freq_sorted(words)
    print "\nFrequency of words:"
    for freq in freq_sorted:
        print freq[0], freq[1]

    output_csv(freq_sorted, "frequency_words.csv")     # open frequency_d3.html to view
    render_histogram_matplot(freq_sorted)
    """
    """
    Finally can put this into a trie if you wanted to search this
    """
